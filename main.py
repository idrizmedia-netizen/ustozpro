import streamlit as st
import sqlite3
import pandas as pd
from docx import Document
import io
from datetime import datetime
import google.generativeai as genai

# --- 0. GEMINI AI SOZLAMASI (ENG BARQAROR VERSIYA) ---
GEMINI_API_KEY = "AIzaSyAox2XPBv1WoKQwBi1K_V8-6VwFssWDyGU"
genai.configure(api_key=GEMINI_API_KEY)

# Bu yerda modelni 'gemini-pro' ga o'zgartirdik
try:
    model = genai.GenerativeModel('gemini-pro')
except Exception as e:
    st.error(f"Modelni yuklashda xatolik: {e}")

# --- 1. MA'LUMOTLAR BAZASI FUNKSIYALARI ---
def create_db():
    conn = sqlite3.connect('ustoz_pro.db', check_same_thread=False)
    c = conn.cursor()
    c.execute('''CREATE TABLE IF NOT EXISTS students 
                 (id INTEGER PRIMARY KEY AUTOINCREMENT, name TEXT, grade TEXT)''')
    c.execute('''CREATE TABLE IF NOT EXISTS marks 
                 (id INTEGER PRIMARY KEY AUTOINCREMENT, student_id INTEGER, 
                  mark TEXT, comment TEXT, date TEXT)''')
    conn.commit()
    return conn

conn = create_db()

# --- 2. AI VA DOCX FUNKSIYALARI ---

def generate_ai_content(mavzu):
    # Promptni sodda va aniq qildik
    full_prompt = f"O'qituvchi uchun '{mavzu}' mavzusida o'zbek tilida batafsil dars konspekti tayyorlab ber."
    
    try:
        response = model.generate_content(full_prompt)
        if response.text:
            return response.text
        else:
            return "AI javob qaytara olmadi. Mavzuni boshqacha yozib ko'ring."
    except Exception as e:
        return f"AI xatoligi: {str(e)}"

def create_docx(mavzu, content):
    doc = Document()
    # ... (bu yerda create_docx funksiyasining qolgan kodi turadi)
    doc.add_heading('Ustoz Pro | AI Dars Konspekti', 0)
    doc.add_paragraph(f"Sana: {datetime.now().strftime('%Y-%m-%d')}")
    doc.add_paragraph(f"Mavzu: {mavzu}")
    
    # AI dan kelgan matnni Word'ga joylash
    doc.add_paragraph(content)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# --- 3. GOOGLE AUTH (KIRISH QISMI) ---
if 'logged_in' not in st.session_state:
    st.session_state.logged_in = False

def login_page():
    st.title("🚀 Ustoz Pro — Kirish")
    st.info("Tizimdan foydalanish uchun shaxsingizni tasdiqlang.")
    if st.button("🔴 Google orqali kirish"):
        st.session_state.logged_in = True
        st.session_state.user_name = "Normurodov I.B." 
        st.rerun()

if not st.session_state.logged_in:
    login_page()
    st.stop()

# --- 4. ASOSIY SAHIFA SOZLAMALARI ---
st.set_page_config(page_title="Ustoz Pro | AI Digital System", layout="wide")

col_head1, col_head2 = st.columns([3, 1])
with col_head1:
    st.subheader(f"👨‍🏫 O'qituvchi: {st.session_state.user_name}")
with col_head2:
    if st.button("Tizimdan chiqish"):
        st.session_state.logged_in = False
        st.rerun()

st.divider()

# --- 5. SIDEBAR (MENYU) ---
st.sidebar.image("https://cdn-icons-png.flaticon.com/512/3426/3426653.png", width=80)
st.sidebar.title("Ustoz Pro Menyu")
page = st.sidebar.radio("Bo'limni tanlang:", ["Jurnal", "O'quvchi qo'shish", "AI Konspekt Generator"])

# --- 6. BO'LIMLAR ---

# A. JURNAL BO'LIMI
if page == "Jurnal":
    st.header("📖 Sinf Jurnali")
    students_df = pd.read_sql_query("SELECT * FROM students", conn)
    
    if not students_df.empty:
        sinflar = students_df['grade'].unique()
        tanlangan_sinf = st.selectbox("📂 Sinfni tanlang:", sinflar)
        st.markdown(f"#### {tanlangan_sinf} sinfi jurnali — {datetime.now().strftime('%d.%m.%Y')}")
        
        filtrlangan = students_df[students_df['grade'] == tanlangan_sinf]
        h_col1, h_col2, h_col3, h_col4 = st.columns([3, 1, 3, 1])
        h_col1.write("**F.I.Sh**")
        h_col2.write("**Baho**")
        h_col3.write("**Izoh**")
        h_col4.write("**Amal**")
        st.divider()

        for index, row in filtrlangan.iterrows():
            c1, c2, c3, c4 = st.columns([3, 1, 3, 1])
            with c1:
                st.write(f"{index+1}. {row['name']}")
            with c2:
                st.selectbox("Baho", ["-", "5", "4", "3", "2"], key=f"m_{row['id']}", label_visibility="collapsed")
            with c3:
                st.text_input("Izoh", placeholder="Darsda faol...", key=f"c_{row['id']}", label_visibility="collapsed")
            with c4:
                if st.button("Saqlash", key=f"b_{row['id']}"):
                    st.toast(f"{row['name']} baholandi!", icon="✅")
    else:
        st.warning("O'quvchilar topilmadi. Avval o'quvchi qo'shing.")

# B. O'QUVCHI QO'SHISH BO'LIMI
elif page == "O'quvchi qo'shish":
    st.header("🆕 O'quvchilarni tizimga kiritish")
    tab1, tab2 = st.tabs(["Yakkalik qo'shish", "Excel/CSV orqali yuklash"])
    
    with tab1:
        col1, col2 = st.columns(2)
        with col1:
            ism = st.text_input("O'quvchi F.I.Sh:", key="single_name")
        with col2:
            sinf = st.text_input("Sinf (masalan: 9-A):", key="single_grade")
        if st.button("Bazaga saqlash", key="save_single"):
            if ism and sinf:
                cursor = conn.cursor()
                cursor.execute("INSERT INTO students (name, grade) VALUES (?, ?)", (ism, sinf))
                conn.commit()
                st.success(f"{ism} muvaffaqiyatli qo'shildi!")
                st.balloons()

    with tab2:
        st.subheader("📁 Fayl orqali ommaviy yuklash")
        uploaded_file = st.file_uploader("Excel yoki CSV faylni tanlang", type=["xlsx", "csv"])
        if uploaded_file is not None:
            try:
                df = pd.read_csv(uploaded_file) if uploaded_file.name.endswith('.csv') else pd.read_excel(uploaded_file)
                st.write("Fayl tarkibi:", df.head())
                if st.button("Bazaga yuklash"):
                    cursor = conn.cursor()
                    for _, row in df.iterrows():
                        cursor.execute("INSERT INTO students (name, grade) VALUES (?, ?)", (str(row['name']), str(row['grade'])))
                    conn.commit()
                    st.success("Muvaffaqiyatli yuklandi!")
            except Exception as e:
                st.error(f"Xatolik: {e}")

# C. AI KONSEPT GENERATOR BO'LIMI
elif page == "AI Konspekt Generator":
    st.header("🤖 Gemini AI Konspekt Generator")
    mavzu = st.text_input("Dars mavzusini kiriting:", placeholder="Masalan: Optika yoki Mexanik harakat")
    
    if st.button("AI yordamida yaratish"):
        if mavzu:
            with st.spinner('Gemini AI dars ishlanmasini tayyorlamoqda...'):
                try:
                    # Modelni chaqirishda xatolik ehtimolini kamaytirish
                    ai_content = generate_ai_content(mavzu)
                    
                    st.markdown("### 📝 AI tomonidan tayyorlangan konspekt:")
                    st.info(ai_content) # Matnni chiroyli ko'rinishda chiqarish
                    
                    file_data = create_docx(mavzu, ai_content)
                    st.download_button(
                        label="📄 Word faylni yuklab olish",
                        data=file_data,
                        file_name=f"{mavzu}_AI_konspekt.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
                except Exception as e:
                    st.error(f"AI bilan bog'lanishda xatolik: {e}")
                    st.info("Maslahat: API kalitini yangilab ko'ring yoki model nomini 'gemini-pro' ga o'zgartiring.")
        else:
            st.warning("Iltimos, mavzu nomini kiriting!")
